from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = Path(__file__).with_name("matriz-acceso-roles-combinados.docx")

NAVY = "17365D"
BLUE = "2E74B5"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D9E0E8"
MUTED = "5B6573"
GREEN = "1E6B43"
AMBER = "8A5A00"
RED = "9B1C1C"
WHITE = "FFFFFF"


DEST_ROLES = [
    "Coordinador de Destacamento",
    "Coordinador Asistente de Destacamento",
    "Pastor de Destacamento",
    "Consejo de Destacamento",
    "Capellán de Destacamento",
    "Líder de Grupo",
    "Líder Asistente de Grupo",
]

SECTION_ROLES = [
    "Coordinador Seccional",
    "Subcoordinador Seccional",
    "Coordinador Seccional de Adiestramiento",
    "Coordinador Seccional de Promoción",
    "Coordinador Seccional de Producción",
    "Coordinador Seccional de Programa",
    "Capellán Seccional",
    "Zonas",
    "Grupos Locales",
]

REGION_ROLES = [
    "Coordinador Regional",
    "Subdirector Regional",
    "Coordinador Regional de Adiestramiento",
    "Coordinador Regional de Promoción",
    "Coordinador Regional de Producción",
    "Coordinador Regional de Programa",
    "Capellán Regional",
    "Secretario Regional",
]

NATIONAL_ROLES = [
    "Consejo Nacional",
    "Ministerios Infantiles Nacional",
    "Director Nacional",
    "Capellán Nacional",
    "Coordinador Nacional de Adiestramiento",
    "Subdirector Nacional",
    "Coordinador Nacional de Promoción",
    "Coordinador Nacional de Producción",
    "Coordinador Nacional de Programa",
    "Comités Especiales Nacional",
    "Oficiales de Adiestramientos Especiales Nacional",
    "Consejo Ejecutivo",
    "Oficina Nacional",
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths):
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            set_cell_margins(cell)


def set_font(run, name="Calibri", size=11, color="000000", bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_font(run, size=9, color=MUTED)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    specs = {
        "Title": (28, NAVY, 0, 8),
        "Subtitle": (13, MUTED, 0, 16),
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, NAVY, 8, 4),
    }
    for name, (size, color, before, after) in specs.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    hp = section.header.paragraphs[0]
    hp.text = "Matriz de acceso organizacional"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_after = Pt(0)
    set_font(hp.runs[0], size=9, color=MUTED, bold=True)

    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    r = fp.add_run("Página ")
    set_font(r, size=9, color=MUTED)
    add_field(fp, " PAGE ")


def add_callout(doc, title, body, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_font(r, size=11, color=NAVY, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_font(r2, size=10.5, color="243447")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def access_color(text):
    if text.startswith("SÍ"):
        return GREEN
    if text.startswith("NO"):
        return RED
    return AMBER


def add_matrix(doc, roles, rule, title):
    doc.add_heading(title, level=2)
    intro = doc.add_paragraph(
        "Cada fila representa siete combinaciones: el rol superior indicado combinado con cada uno de los siete cargos de destacamento definidos en la sección anterior."
    )
    intro.paragraph_format.keep_with_next = True

    headers = ["Rol superior", "Combinado con", "Regiones", "Secciones", "Destacamentos", "Miembros"]
    widths = [1800, 1800, 1440, 1440, 1440, 1440]
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, label in enumerate(headers):
        cell = header.cells[idx]
        set_cell_shading(cell, NAVY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        set_font(r, size=8.5, color=WHITE, bold=True)

    for role in roles:
        row = table.add_row()
        set_row_cant_split(row)
        values = [role, "Cada cargo de destacamento", *rule]
        for idx, value in enumerate(values):
            cell = row.cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx < 2 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(value)
            set_font(r, size=8.2, color=access_color(value) if idx >= 2 else "1F2937", bold=idx == 0)
    set_table_geometry(table, widths)


def add_dest_roles_table(doc):
    doc.add_heading("Cargos de destacamento incluidos", level=1)
    p = doc.add_paragraph(
        "Estos siete cargos son la base de todas las combinaciones del anexo. El cargo local conserva sus facultades dentro de su destacamento; el cargo superior amplía la visibilidad según su nivel."
    )
    p.paragraph_format.keep_with_next = True

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    headers = ["Código", "Cargo de destacamento"]
    for i, label in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], NAVY)
        p = table.rows[0].cells[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        set_font(r, size=9, color=WHITE, bold=True)
    set_repeat_table_header(table.rows[0])
    for idx, role in enumerate(DEST_ROLES, 1):
        row = table.add_row()
        set_row_cant_split(row)
        for col, value in enumerate((f"D{idx}", role)):
            p = row.cells[col].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_font(r, size=9.5, color="1F2937", bold=col == 0)
    set_table_geometry(table, [1200, 8160])


def add_base_summary(doc):
    doc.add_heading("Resumen de alcance sin combinación superior", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["Perfil", "Regiones", "Secciones", "Destacamentos", "Miembros"]
    widths = [2160, 1800, 1800, 1800, 1800]
    for idx, label in enumerate(headers):
        set_cell_shading(table.rows[0].cells[idx], NAVY)
        p = table.rows[0].cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        set_font(r, size=8.5, color=WHITE, bold=True)
    set_repeat_table_header(table.rows[0])

    rows = [
        ("Usuario común", "Todas listadas; abre la propia", "Su región", "Su sección", "Su destacamento"),
        ("Cualquier cargo D1–D7", "Todas listadas; abre la propia", "Su región", "Su región", "Su destacamento"),
    ]
    for values in rows:
        row = table.add_row()
        set_row_cant_split(row)
        for idx, value in enumerate(values):
            p = row.cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_font(r, size=8.5, color="1F2937", bold=idx == 0)
    set_table_geometry(table, widths)


def build():
    doc = Document()
    configure_styles(doc)
    for section in doc.sections:
        configure_section(section)

    # Editorial-cover opening, adapted to a compact internal reference guide.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(72)
    p.paragraph_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("GUÍA DE REFERENCIA")
    set_font(r, size=11, color=BLUE, bold=True)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Matriz de acceso para roles combinados")

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(
        "Regiones, secciones, destacamentos y miembros\nCombinaciones con cargos de destacamento"
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Estado del sistema revisado el 29 de agosto de 2026")
    set_font(r, size=10.5, color=MUTED, italic=True)

    doc.add_page_break()

    doc.add_heading("Cómo leer este documento", level=1)
    doc.add_paragraph(
        "La palabra “ver” se refiere a aparecer en la lista y poder consultar el nivel indicado. En Regiones existe una distinción: los cargos inferiores ven todas las regiones listadas, pero solamente pueden abrir la propia; por eso no se considera acceso nacional completo."
    )
    add_callout(
        doc,
        "Supuesto organizativo",
        "Las combinaciones se evalúan dentro de la misma cadena organizativa: el cargo de destacamento pertenece a la sección y región del cargo superior. Si se asignan cargos en regiones distintas, el alcance debe revisarse como unión de asignaciones y no como una sola cadena.",
    )
    add_callout(
        doc,
        "Qué no cubre la matriz",
        "La matriz describe visibilidad de listas. No concede edición, eliminación, acceso a datos sensibles ni aprobación. Esas capacidades se determinan por permisos adicionales y por la entidad propia de cada cargo.",
        fill=LIGHT_GRAY,
    )

    add_dest_roles_table(doc)
    add_base_summary(doc)

    doc.add_page_break()
    doc.add_heading("Combinaciones exhaustivas por nivel", level=1)
    doc.add_paragraph(
        "En total se documentan 210 combinaciones: 30 roles superiores multiplicados por siete cargos de destacamento. Como los siete cargos locales producen la misma visibilidad cuando se combinan con un mismo nivel superior, cada fila agrupa esas siete combinaciones sin perder cobertura."
    )

    section_rule = (
        "NO — listadas; abre su región",
        "NO — su región",
        "NO — su región",
        "NO — su sección",
    )
    regional_rule = (
        "SÍ — todas",
        "SÍ — todas",
        "NO — su región",
        "NO — su región",
    )
    national_rule = (
        "SÍ — todas",
        "SÍ — todas",
        "SÍ — todos",
        "SÍ — todos",
    )

    add_matrix(doc, SECTION_ROLES, section_rule, "Nivel seccional × cargos de destacamento")
    add_matrix(doc, REGION_ROLES, regional_rule, "Nivel regional × cargos de destacamento")
    add_matrix(doc, NATIONAL_ROLES, national_rule, "Consejo Nacional/Ejecutivo × cargos de destacamento")

    doc.add_heading("Conclusiones", level=1)
    doc.add_paragraph(
        "Una combinación seccional no obtiene visibilidad nacional: conserva las secciones y destacamentos de su región y los miembros de su sección. Una combinación regional ve todas las regiones y secciones, pero los destacamentos y miembros permanecen en su región. Una combinación nacional o ejecutiva ve todos los registros de los cuatro niveles."
    )
    doc.add_paragraph(
        "El cargo de destacamento nunca pierde su autoridad local por combinarse con un cargo superior; lo que cambia en esta matriz es la amplitud de consulta."
    )

    doc.core_properties.title = "Matriz de acceso para roles combinados"
    doc.core_properties.subject = "Visibilidad de regiones, secciones, destacamentos y miembros"
    doc.core_properties.author = "Sistema de Exploradores"
    doc.core_properties.keywords = "roles, permisos, alcance, regiones, secciones, destacamentos, miembros"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
