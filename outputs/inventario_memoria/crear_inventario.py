from __future__ import annotations

import os
import sys
from collections import defaultdict
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE = Path(r"D:\documentos\Documentos Roderi\Exploradores\Memoria 2.0")
ROOT_NAMES = ("Exploradores", "Seguidores", "Pioneros")
OUTPUT = BASE / "Inventario de archivos - Exploradores, Seguidores y Pioneros.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "263238"
MUTED = "66717A"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "B8C4CE"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(run, size=None, bold=None, color=INK, italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = rgb(color)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_field(paragraph, instruction: str):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=8.5, color=MUTED)


def collect_inventory():
    inventory = {}
    for root_name in ROOT_NAMES:
        root = BASE / root_name
        if not root.is_dir():
            raise FileNotFoundError(f"No existe la carpeta requerida: {root}")
        grouped = defaultdict(list)
        for current_dir, dir_names, file_names in os.walk(root):
            dir_names.sort(key=str.casefold)
            current = Path(current_dir)
            relative_dir = current.relative_to(root)
            key = "." if str(relative_dir) == "." else str(relative_dir)
            for file_name in sorted(file_names, key=str.casefold):
                grouped[key].append(file_name)
        inventory[root_name] = dict(sorted(grouped.items(), key=lambda item: item[0].casefold()))
    return inventory


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.widow_control = True

    for style_name, size, color, before, after in (
        ("Title", 24, DARK_BLUE, 0, 5),
        ("Subtitle", 11.5, MUTED, 0, 16),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 12.5, DARK_BLUE, 10, 5),
        ("Heading 3", 11, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
        style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = style_name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True

    bullet = doc.styles["List Bullet"]
    bullet.font.name = "Calibri"
    bullet._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    bullet._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    bullet.font.size = Pt(9.5)
    bullet.font.color.rgb = rgb(INK)
    bullet.paragraph_format.left_indent = Inches(0.375)
    bullet.paragraph_format.first_line_indent = Inches(-0.188)
    bullet.paragraph_format.space_after = Pt(2)
    bullet.paragraph_format.line_spacing = 1.05
    bullet.paragraph_format.widow_control = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    run = hp.add_run("INVENTARIO DOCUMENTAL  |  MEMORIA 2.0")
    set_run_font(run, size=8.5, bold=True, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    label = fp.add_run("Página ")
    set_run_font(label, size=8.5, color=MUTED)
    add_field(fp, "PAGE")


def add_summary(doc: Document, inventory):
    counts = {
        root: sum(len(files) for files in groups.values())
        for root, groups in inventory.items()
    }
    total = sum(counts.values())

    title = doc.add_paragraph(style="Title")
    title.add_run("Inventario de archivos")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("Carpetas Exploradores, Seguidores y Pioneros")

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(12)
    r = meta.add_run("Ubicación revisada: ")
    set_run_font(r, size=9.5, bold=True, color=MUTED)
    r = meta.add_run(str(BASE))
    set_run_font(r, size=9.5, color=MUTED)
    r = meta.add_run(f"\nFecha del inventario: {datetime.now().strftime('%d/%m/%Y')}  |  Total: {total} archivos")
    set_run_font(r, size=9.5, color=MUTED)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.autofit = False
    headers = ("Carpeta", "Subcarpetas con archivos", "Archivos")
    for index, value in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(value)
        set_run_font(run, size=9.5, bold=True, color=DARK_BLUE)

    for root_name in ROOT_NAMES:
        row = table.add_row()
        values = (root_name, str(len(inventory[root_name])), str(counts[root_name]))
        for index, value in enumerate(values):
            cell = row.cells[index]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            set_run_font(run, size=9.5, bold=index == 0, color=INK)

    total_row = table.add_row()
    for cell in total_row.cells:
        set_cell_shading(cell, LIGHT_GRAY)
    values = ("Total", str(sum(len(groups) for groups in inventory.values())), str(total))
    for index, value in enumerate(values):
        p = total_row.cells[index].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(value)
        set_run_font(run, size=9.5, bold=True, color=DARK_BLUE)

    set_table_geometry(table, [4300, 2700, 1760])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_inventory_sections(doc: Document, inventory):
    for root_index, root_name in enumerate(ROOT_NAMES):
        doc.add_page_break()
        root_heading = doc.add_paragraph(style="Heading 1")
        root_heading.add_run(root_name)
        count = sum(len(files) for files in inventory[root_name].values())
        info = doc.add_paragraph()
        info.paragraph_format.space_after = Pt(8)
        run = info.add_run(f"{count} archivos organizados por ubicación relativa.")
        set_run_font(run, size=9.5, italic=True, color=MUTED)

        groups = inventory[root_name]
        if not groups:
            p = doc.add_paragraph()
            p.add_run("No se encontraron archivos.")
            continue

        for relative_dir, files in groups.items():
            sub_heading = doc.add_paragraph(style="Heading 2")
            display = "Carpeta principal" if relative_dir == "." else relative_dir
            sub_heading.add_run(display)
            count_run = sub_heading.add_run(f"  ({len(files)})")
            set_run_font(count_run, size=10, bold=False, color=MUTED)

            for file_name in files:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(file_name)


def main():
    inventory = collect_inventory()
    doc = Document()
    configure_document(doc)
    add_summary(doc, inventory)
    add_inventory_sections(doc, inventory)
    doc.core_properties.title = "Inventario de archivos - Memoria 2.0"
    doc.core_properties.subject = "Listado de archivos de Exploradores, Seguidores y Pioneros"
    doc.core_properties.author = ""
    doc.core_properties.keywords = "inventario, archivos, memoria"
    doc.save(OUTPUT)
    print(str(OUTPUT))
    for root_name, groups in inventory.items():
        print(f"{root_name}: {sum(len(files) for files in groups.values())}")


if __name__ == "__main__":
    main()
