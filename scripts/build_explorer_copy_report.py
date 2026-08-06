#!/usr/bin/env python3
"""Build the comprehensive Explorer-themed UI copy audit report."""

from __future__ import annotations

import argparse
import json
import re
import sys
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


NAVY = RGBColor(32, 55, 72)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(92, 105, 117)
GOLD = RGBColor(174, 121, 22)
INK = RGBColor(25, 36, 45)
WHITE = RGBColor(255, 255, 255)
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"
GREEN_FILL = "E7F4EC"
GOLD_FILL = "FFF4D6"
RED_FILL = "FCE8E6"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 100
CELL_MARGINS = {"top": 65, "bottom": 65, "start": 100, "end": 100}


def set_run_font(run, size=11, color=INK, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tag = OxmlElement("w:tblHeader")
    tag.set(qn("w:val"), "true")
    tr_pr.append(tag)


def prevent_split(row):
    row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    set_run_font(run, size=8.5, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def add_bottom_border(paragraph, color="D7DBE2"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p_pr.append(borders)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.widow_control = True

    specs = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (11.5, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in specs.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.18

    header = section.header.paragraphs[0]
    header.text = ""
    header.paragraph_format.space_after = Pt(3)
    set_run_font(
        header.add_run("EDR DASHBOARD 1.1  |  AUDITORÍA DE LENGUAJE EXPLORADORÍSTICO"),
        size=8.2,
        color=MUTED,
        bold=True,
    )
    add_bottom_border(header)
    footer = section.footer.paragraphs[0]
    footer.text = ""
    add_page_field(footer)


def body(doc, text, *, italic=False):
    paragraph = doc.add_paragraph()
    set_run_font(paragraph.add_run(text), size=10.5, italic=italic)
    return paragraph


def bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    set_run_font(paragraph.add_run(text), size=10.3)
    return paragraph


def callout(doc, title, text, apply_table_geometry, fill=LIGHT_FILL):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    repeat_header(table.rows[0])
    prevent_split(table.rows[0])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    cell.text = ""
    first = cell.paragraphs[0]
    first.paragraph_format.space_after = Pt(3)
    set_run_font(first.add_run(title), size=10.2, color=NAVY, bold=True)
    second = cell.add_paragraph()
    second.paragraph_format.space_after = Pt(0)
    set_run_font(second.add_run(text), size=9.6)
    apply_table_geometry(
        table,
        [TABLE_WIDTH_DXA],
        table_width_dxa=TABLE_WIDTH_DXA,
        indent_dxa=TABLE_INDENT_DXA,
        cell_margins_dxa=CELL_MARGINS,
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def safe_lower_first(text):
    if not text:
        return text
    return text[0].lower() + text[1:]


def clean_dynamic(text):
    return text.replace("{…}", "el elemento").replace("{...}", "el elemento").strip()


EXACT_OPTIONS = {
    "iniciar sesión": ["Entrar al campamento", "Abrir mi bitácora", "Iniciar mi jornada"],
    "crear cuenta": ["Unirme a la expedición", "Crear mi bitácora", "Comenzar como Explorador"],
    "registrarse": ["Unirme a la expedición", "Registrar mi ingreso", "Preparar mi bitácora"],
    "sign in": ["Entrar al campamento", "Abrir mi bitácora", "Iniciar mi jornada"],
    "create account": ["Unirme a la expedición", "Crear mi bitácora", "Comenzar como Explorador"],
    "guardar": ["Guardar en la bitácora", "Registrar este avance", "Confirmar el registro"],
    "guardar cambios": ["Guardar en la bitácora", "Registrar los cambios", "Confirmar este avance"],
    "cancelar": ["Volver al sendero", "Dejar sin cambios", "Cancelar la acción"],
    "cerrar": ["Cerrar la bitácora", "Volver al campamento", "Cerrar"],
    "buscar": ["Explorar registros", "Buscar en el mapa", "Seguir una pista"],
    "eliminar": ["Retirar del registro", "Eliminar de la bitácora", "Confirmar retiro"],
    "actualizar": ["Actualizar la bitácora", "Registrar el nuevo estado", "Renovar la información"],
    "crear": ["Iniciar una nueva misión", "Crear registro", "Agregar a la bitácora"],
    "enviar": ["Enviar señal", "Compartir con la patrulla", "Enviar mensaje"],
    "volver": ["Volver al sendero", "Regresar al campamento", "Volver"],
    "continuar": ["Seguir la ruta", "Continuar la jornada", "Avanzar"],
    "add card": ["Agregar tarjeta", "Crear una tarea", "Registrar una nueva misión"],
    "add folder": ["Agregar carpeta", "Crear carpeta de expedición", "Abrir una nueva carpeta"],
    "on upload": ["Al subir el archivo", "Cuando se complete la carga", "Después de guardar el documento"],
    "can edit": ["Puede editar", "Puede actualizar el registro", "Tiene permiso de edición"],
    "delete success!": ["Elemento eliminado correctamente", "Registro retirado de la bitácora", "Misión cumplida: elemento eliminado"],
}


def alternatives(text, surface, priority):
    cleaned = clean_dynamic(text)
    lowered = cleaned.lower().strip(" .:;¡!¿?")
    if lowered in EXACT_OPTIONS:
        return EXACT_OPTIONS[lowered]

    if re.fullmatch(r"(?:iniciar sesión|sign in|entrar|entrar como .+|acceder|acceso)", lowered):
        return [
            "Entrar al campamento",
            "Abrir mi bitácora de Explorador",
            "Iniciar la jornada en el sistema",
        ]
    if re.search(r"crear cuenta|registr|sign up|crear usuario", lowered):
        return [
            "Unirme a la expedición",
            "Crear mi registro de Explorador",
            "Preparar mi acceso al campamento",
        ]
    if re.search(r"contraseña|password|recuper", lowered):
        return [
            f"Recuperar el acceso · {cleaned}",
            f"Volver a la ruta · {cleaned}",
            f"Restablecer la clave de mi bitácora · {cleaned}",
        ]
    if re.search(r"cargando|loading|preparando|procesando", lowered):
        return [
            "Preparando el campamento…",
            "Consultando la bitácora…",
            "Trazando la ruta…",
        ]
    if re.search(r"sin resultados|no se encontr|no hay|vacío|empty|aún no", lowered):
        return [
            f"Aún no encontramos huellas. {cleaned}",
            f"La bitácora está en blanco por ahora. {cleaned}",
            f"No hay registros en esta ruta. {cleaned}",
        ]
    if re.search(r"éxito|correctamente|completad|guardad|cread|actualizad|enviad|aprobado", lowered):
        return [
            f"Misión cumplida: {safe_lower_first(cleaned)}",
            f"Avance registrado en la bitácora: {safe_lower_first(cleaned)}",
            f"La patrulla puede continuar: {safe_lower_first(cleaned)}",
        ]
    if re.search(r"permiso|autoriz|acceso denegado|no tienes", lowered):
        return [
            f"Esta ruta requiere autorización. {cleaned}",
            f"Consulta con tu líder para continuar. {cleaned}",
            f"Acceso protegido del campamento. {cleaned}",
        ]
    if re.search(r"error|fall|inválid|no se pudo|expir|obligatori|requerid", lowered) or surface == "Error o validación":
        return [
            f"Encontramos un obstáculo en la ruta. {cleaned}",
            f"No pudimos completar esta misión. {cleaned}",
            f"Revisa la bitácora e inténtalo otra vez. {cleaned}",
        ]
    if re.search(r"notific|nuevo mensaje|recib|aviso", lowered) or surface == "Notificación":
        return [
            f"Nueva señal en el campamento: {safe_lower_first(cleaned)}",
            f"La bitácora tiene novedades: {safe_lower_first(cleaned)}",
            f"Atención, Explorador: {safe_lower_first(cleaned)}",
        ]
    if re.search(r"buscar|resultado|filtro|explorar", lowered):
        return [
            f"Seguir una pista · {cleaned}",
            f"Explorar la bitácora · {cleaned}",
            f"Buscar en el mapa · {cleaned}",
        ]
    if re.search(r"guardar|actualizar|editar|cambiar", lowered):
        return [
            f"Registrar este avance · {cleaned}",
            f"Actualizar la bitácora · {cleaned}",
            f"Confirmar los cambios · {cleaned}",
        ]
    if re.search(r"crear|agregar|añadir|nuevo|subir|upload", lowered):
        return [
            f"Iniciar una nueva misión · {cleaned}",
            f"Agregar a la bitácora · {cleaned}",
            f"Preparar un nuevo registro · {cleaned}",
        ]
    if re.search(r"enviar|mensaje|chat|conversación|responder", lowered):
        return [
            f"Enviar una señal · {cleaned}",
            f"Compartir con la patrulla · {cleaned}",
            f"Continuar la conversación · {cleaned}",
        ]
    if re.search(r"eliminar|borrar|quitar|retirar|limpiar", lowered):
        return [
            f"Retirar del registro · {cleaned}",
            f"Eliminar de la bitácora · {cleaned}",
            f"Confirmar el retiro · {cleaned}",
        ]
    if re.search(r"confirm|seguro|advertencia|atención", lowered) or surface == "Aviso o confirmación":
        return [
            f"Confirma el siguiente paso de la misión. {cleaned}",
            f"Antes de continuar la ruta: {safe_lower_first(cleaned)}",
            f"Revisa esta decisión en la bitácora. {cleaned}",
        ]
    if re.search(r"perfil|cuenta|miembro|usuario", lowered):
        return [
            f"Bitácora del Explorador · {cleaned}",
            f"Registro del miembro · {cleaned}",
            f"Datos de la patrulla · {cleaned}",
        ]
    if surface == "Botón o acción":
        return [
            f"{cleaned} en la bitácora",
            f"{cleaned} para continuar la ruta",
            cleaned,
        ]
    if priority.startswith("Baja"):
        return [
            cleaned,
            f"{cleaned} Contacta a tu líder si necesitas ayuda.",
            f"Acción protegida: {safe_lower_first(cleaned)}",
        ]
    return [
        f"Bitácora · {cleaned}",
        f"Ruta de Exploradores · {cleaned}",
        f"Campamento · {cleaned}",
    ]


def has_english(text):
    return bool(re.search(r"\b(add|create|delete|edit|empty|folder|loading|password|request|save|sign|update|upload|welcome|with)\b", text, re.I))


def score_item(item):
    score = 0
    module = item["module"]
    surface = item["surface"]
    text = item["text"]
    score += {
        "Autenticación y acceso": 60,
        "Chat y mensajería": 55,
        "Miembros": 35,
        "Administración y permisos": 32,
        "Navegación y estructura": 28,
        "Componentes compartidos": 25,
    }.get(module, 18)
    score += {
        "Notificación": 42,
        "Aviso o confirmación": 38,
        "Error o validación": 36,
        "Botón o acción": 32,
        "Campo o ayuda": 28,
        "Navegación": 25,
        "Título o contenido": 15,
        "Texto general": 5,
    }.get(surface, 0)
    score += {"Alta": 30, "Media": 15, "Baja: conservar claridad": -5}.get(item["priority"], 0)
    if has_english(text):
        score += 25
    if re.search(r"bienven|iniciar sesión|crear cuenta|sin resultados|cargando|correctamente|no se pudo|eliminar|guardar|buscar|enviar", text, re.I):
        score += 20
    if re.search(r"^\[|console|test\]|ERROR DETALLADO|👉|&:|checkbox$", text, re.I):
        score -= 40
    return score


def select_catalog(items):
    by_module = defaultdict(dict)
    for item in items:
        key = item["text"].lower()
        existing = by_module[item["module"]].get(key)
        if existing is None:
            by_module[item["module"]][key] = {
                **item,
                "surfaces": {item["surface"]},
                "sources": list(item["sources"]),
            }
        else:
            existing["surfaces"].add(item["surface"])
            for source in item["sources"]:
                if source not in existing["sources"] and len(existing["sources"]) < 8:
                    existing["sources"].append(source)

    caps = defaultdict(lambda: 6)
    caps.update(
        {
            "Autenticación y acceso": 30,
            "Chat y mensajería": 22,
            "Miembros": 20,
            "Administración y permisos": 15,
            "Componentes compartidos": 14,
            "Páginas y metadatos": 12,
            "Navegación y estructura": 12,
            "API y mensajes del servidor": 12,
            "Servicios y operaciones": 10,
            "Archivos": 10,
            "Certificados": 10,
            "Usuarios": 10,
        }
    )

    selected = []
    for module, module_items in by_module.items():
        ranked = sorted(module_items.values(), key=lambda item: (-score_item(item), item["text"]))
        selected.extend(ranked[: caps[module]])
    return sorted(selected, key=lambda item: (item["module"], -score_item(item), item["text"]))


def set_cell_paragraph(cell, text, *, bold=False, size=8.5, color=INK, italic=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(1.5)
    paragraph.paragraph_format.line_spacing = 1.05
    set_run_font(paragraph.add_run(str(text)), size=size, color=color, bold=bold, italic=italic)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    return paragraph


def add_catalog_table(doc, entries, apply_table_geometry):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    repeat_header(table.rows[0])
    prevent_split(table.rows[0])
    for index, text in enumerate(("Texto actual y ubicación", "Alternativas exploradorísticas")):
        shade_cell(table.rows[0].cells[index], HEADER_FILL)
        set_cell_paragraph(table.rows[0].cells[index], text, bold=True, size=8.8, color=NAVY)

    for row_index, entry in enumerate(entries):
        row = table.add_row()
        prevent_split(row)
        if row_index % 2:
            shade_cell(row.cells[0], "FAFBFC")
            shade_cell(row.cells[1], "FAFBFC")
        current = row.cells[0]
        paragraph = set_cell_paragraph(current, f'“{entry["text"]}”', bold=True, size=8.3)
        meta = current.add_paragraph()
        meta.paragraph_format.space_after = Pt(0)
        meta.paragraph_format.line_spacing = 1.0
        surfaces = ", ".join(sorted(entry.get("surfaces", {entry["surface"]})))
        source = entry["sources"][0]
        set_run_font(
            meta.add_run(f"{surfaces} · {entry['priority']}\n{source}"),
            size=7.2,
            color=MUTED,
            italic=True,
        )

        alt_cell = row.cells[1]
        alt_cell.text = ""
        options = alternatives(entry["text"], entry["surface"], entry["priority"])
        for index, option in enumerate(options, start=1):
            p = alt_cell.paragraphs[0] if index == 1 else alt_cell.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.02
            set_run_font(p.add_run(f"{index}. {option}"), size=8.1)
        alt_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    apply_table_geometry(
        table,
        [3800, 5560],
        table_width_dxa=TABLE_WIDTH_DXA,
        indent_dxa=TABLE_INDENT_DXA,
        cell_margins_dxa=CELL_MARGINS,
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_summary_table(doc, rows, apply_table_geometry):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    repeat_header(table.rows[0])
    prevent_split(table.rows[0])
    for index, text in enumerate(("Módulo", "Textos detectados", "Enfoque recomendado")):
        shade_cell(table.rows[0].cells[index], HEADER_FILL)
        set_cell_paragraph(table.rows[0].cells[index], text, bold=True, size=8.8, color=NAVY)
    for module, count, focus in rows:
        row = table.add_row()
        prevent_split(row)
        set_cell_paragraph(row.cells[0], module, size=8.2)
        set_cell_paragraph(row.cells[1], count, size=8.2)
        set_cell_paragraph(row.cells[2], focus, size=8.2)
    apply_table_geometry(
        table,
        [3200, 1500, 4660],
        table_width_dxa=TABLE_WIDTH_DXA,
        indent_dxa=TABLE_INDENT_DXA,
        cell_margins_dxa=CELL_MARGINS,
    )


def module_focus(module):
    if module == "Autenticación y acceso":
        return "Alta tematización en bienvenida y recuperación; máxima claridad en seguridad."
    if module == "Chat y mensajería":
        return "Tono de señales, patrulla y fogata; acciones destructivas explícitas."
    if module in {"Miembros", "Administración y permisos", "API y mensajes del servidor"}:
        return "Tematización moderada; preservar precisión, permisos y trazabilidad."
    if module in {"Navegación y estructura", "Páginas y metadatos"}:
        return "Unificar nombres de rutas, campamento, bitácora y jornada."
    return "Aplicar tono ligero en estados vacíos, éxito, carga y acciones cotidianas."


def build_report(data_path: Path, output_path: Path, skill_root: Path):
    sys.path.insert(0, str(skill_root / "scripts"))
    from table_geometry import apply_table_geometry

    data = json.loads(data_path.read_text(encoding="utf-8"))
    summary = data["summary"]
    selected = select_catalog(data["items"])
    selected_by_module = defaultdict(list)
    for item in selected:
        selected_by_module[item["module"]].append(item)

    doc = Document()
    configure_document(doc)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(70)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(16)
    set_run_font(kicker.add_run("AUDITORÍA INTEGRAL DE MICROCOPY"), size=10, color=GOLD, bold=True)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(7)
    set_run_font(title.add_run("Lenguaje exploradorístico"), size=29, color=NAVY, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(5)
    set_run_font(
        subtitle.add_run("Mensajes, notificaciones, avisos, botones, login, registro y módulos"),
        size=14.2,
        color=DARK_BLUE,
    )
    descriptor = doc.add_paragraph()
    descriptor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    descriptor.paragraph_format.space_after = Pt(62)
    set_run_font(
        descriptor.add_run("Inventario técnico y propuestas de reemplazo"),
        size=10.2,
        color=GOLD,
        bold=True,
    )
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(meta.add_run("Exploradores del Rey · EDR Dashboard 1.1"), size=12, color=NAVY, bold=True)
    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(date.add_run("6 de agosto de 2026"), size=9.3, color=MUTED, italic=True)
    doc.add_page_break()

    doc.add_heading("1. Resultado del análisis", level=1)
    body(
        doc,
        f"Se inspeccionaron {summary['sourceFiles']:,} archivos de src y se detectaron {summary['uniqueTexts']:,} combinaciones únicas de texto, módulo y superficie. El análisis cubrió mensajes visibles, notificaciones, errores, validaciones, avisos, confirmaciones, botones, campos, navegación, login, registro y respuestas del servidor.",
    )
    body(
        doc,
        f"El catálogo detallado consolida {len(selected):,} puntos de cambio de mayor impacto distribuidos en todos los módulos detectados. Los duplicados exactos y las variantes de una misma frase se agruparon para evitar recomendaciones repetidas. El inventario técnico completo permanece trazable en user-facing-copy.json.",
    )
    callout(
        doc,
        "Criterio rector",
        "El tono exploradorístico debe aportar identidad sin ocultar la acción. En seguridad, permisos, salud, datos sensibles, pagos y eliminaciones, la claridad funcional siempre prevalece sobre la metáfora.",
        apply_table_geometry,
        fill=GREEN_FILL,
    )

    doc.add_heading("2. Sistema de voz recomendado", level=1)
    for text in [
        "Voz: cercana, activa, servicial y orientada a misión; nunca infantil ni excesivamente militar.",
        "Metáforas principales: campamento, sendero, ruta, brújula, bitácora, patrulla, señal, fogata, misión e insignia.",
        "Éxitos: «Misión cumplida» o «Avance registrado»; cargas: «Preparando el campamento»; vacíos: «Aún no encontramos huellas».",
        "Errores: primero explicar qué ocurrió y cómo resolverlo; la metáfora solo introduce el mensaje.",
        "Acciones destructivas: usar verbos literales como eliminar, retirar o borrar y confirmar el alcance.",
        "Accesibilidad: no depender únicamente de emojis, color o humor para comunicar estado o severidad.",
    ]:
        bullet(doc, text)

    doc.add_heading("2.1 Léxico de reemplazo", level=2)
    add_summary_table(
        doc,
        [
            ("Sistema / plataforma", "—", "Campamento, base o centro de operaciones."),
            ("Registro / historial", "—", "Bitácora, registro de ruta o avance."),
            ("Buscar", "—", "Explorar, seguir una pista o buscar en el mapa."),
            ("Mensaje / notificación", "—", "Señal, novedad de la patrulla o aviso del campamento."),
            ("Proceso", "—", "Misión, jornada, ruta o siguiente paso."),
            ("Usuario / miembro", "—", "Explorador o miembro; usar patrulla solo cuando sea grupal."),
        ],
        apply_table_geometry,
    )

    doc.add_heading("3. Hallazgos transversales", level=1)
    callout(
        doc,
        "Textos internos expuestos",
        "Se detectaron frases técnicas como «Token expired!», «Access token not found in response», etiquetas de pruebas y logs con prefijos ERROR. Si alcanzan la interfaz, deben sustituirse por mensajes controlados y registrar el detalle únicamente en observabilidad.",
        apply_table_geometry,
        fill=RED_FILL,
    )
    for text in [
        "Hay mezcla de español e inglés en login, recuperación, carpetas, carga de archivos y plantillas.",
        "Acciones iguales usan variantes distintas: Crear, Crear nuevo, Agregar, Añadir, Guardar y Guardar cambios.",
        "Algunos mensajes técnicos mencionan Firebase, tokens, variables o rutas internas; no deberían mostrarse al usuario final.",
        "Las áreas de permisos, salud y eliminación requieren un tono sobrio, aunque pueden cerrar con una ayuda contextual del líder.",
        "Los estados vacíos, cargas, éxitos, bienvenida y notificaciones son los mejores lugares para reforzar la identidad de Exploradores del Rey.",
    ]:
        bullet(doc, text)

    doc.add_heading("4. Cobertura por módulo", level=1)
    coverage_rows = [
        (module, str(count), module_focus(module))
        for module, count in summary["byModule"].items()
    ]
    add_summary_table(doc, coverage_rows, apply_table_geometry)

    doc.add_page_break()
    doc.add_heading("5. Catálogo de textos y alternativas", level=1)
    body(
        doc,
        "Cada entrada muestra el texto actual, su superficie y ubicación principal. Las tres alternativas representan niveles equivalentes de tematización; la primera se recomienda como punto de partida y debe validarse en contexto antes de reemplazarla.",
    )

    for module in summary["byModule"].keys():
        entries = selected_by_module.get(module, [])
        if not entries:
            continue
        doc.add_heading(module, level=2)
        note = doc.add_paragraph()
        note.paragraph_format.space_after = Pt(4)
        set_run_font(
            note.add_run(
                f"{summary['byModule'][module]} textos detectados · {len(entries)} puntos priorizados en este catálogo."
            ),
            size=8.5,
            color=MUTED,
            italic=True,
        )
        add_catalog_table(doc, entries, apply_table_geometry)

    doc.add_heading("6. Textos que no conviene tematizar en exceso", level=1)
    body(
        doc,
        "Los siguientes tipos pueden incorporar una introducción breve, pero deben conservar el significado literal y la acción correctiva:",
    )
    for text in [
        "Contraseñas, sesión expirada, verificación de identidad y autenticación multifactor.",
        "Permisos, acceso a menores, datos médicos, privacidad y trazabilidad de auditoría.",
        "Eliminación global, retiro de participantes, transferencia de propiedad y pérdida de datos.",
        "Pagos, facturas, montos, estados financieros y confirmaciones de compra.",
        "Errores del servidor: nunca exponer tokens, proveedores, variables de entorno, rutas o trazas internas.",
    ]:
        bullet(doc, text)
    callout(
        doc,
        "Ejemplo seguro",
        "Preferir «Encontramos un obstáculo en la ruta. Tu sesión expiró; inicia sesión nuevamente.» sobre «La brújula se perdió». La primera conserva la identidad de marca y explica exactamente qué hacer.",
        apply_table_geometry,
        fill=GOLD_FILL,
    )

    doc.add_heading("7. Orden recomendado de implementación", level=1)
    rollout = [
        "Fase 1 — Login, registro, recuperación, bienvenida, navegación principal y estados vacíos.",
        "Fase 2 — Chat, notificaciones, éxitos, cargas, búsquedas y acciones frecuentes.",
        "Fase 3 — Miembros, regiones, secciones, destacamentos, archivos y certificados.",
        "Fase 4 — Administración, permisos, salud, eliminaciones, pagos y errores controlados.",
        "Fase 5 — Centralizar textos en un catálogo/i18n, revisar consistencia y ejecutar QA de accesibilidad.",
    ]
    for item in rollout:
        paragraph = doc.add_paragraph(style="List Number")
        set_run_font(paragraph.add_run(item), size=10.3)

    doc.add_heading("8. Criterios de QA de contenido", level=1)
    for text in [
        "La acción sigue siendo comprensible sin conocer la metáfora exploradora.",
        "El texto no infantiliza al usuario ni convierte un error serio en una broma.",
        "Botón, título, aviso y notificación usan el mismo término para la misma acción.",
        "Los mensajes dinámicos conservan nombres, cantidades, fechas y consecuencias.",
        "Cada error indica el siguiente paso o el contacto apropiado.",
        "Lectores de pantalla reciben el mismo significado que usuarios visuales; el emoji es decorativo.",
        "La redacción funciona en móvil sin truncar botones ni diálogos.",
    ]:
        bullet(doc, text)

    body(
        doc,
        "Conclusión: la aplicación posee suficientes puntos de contacto para construir una voz reconocible de Exploradores del Rey. El mayor impacto vendrá de tematizar primero la bienvenida, los estados de avance, las notificaciones y los vacíos; las áreas sensibles deben adoptar una identidad más ligera y conservar su precisión operacional.",
        italic=True,
    )

    doc.core_properties.title = "Auditoría integral de lenguaje exploradorístico"
    doc.core_properties.subject = "Mensajes y alternativas de microcopy por módulo"
    doc.core_properties.keywords = "Exploradores del Rey, microcopy, UX writing, mensajes, notificaciones"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("data", type=Path)
    parser.add_argument("output", type=Path)
    parser.add_argument("--skill-root", type=Path, required=True)
    args = parser.parse_args()
    build_report(args.data, args.output, args.skill_root)


if __name__ == "__main__":
    main()
