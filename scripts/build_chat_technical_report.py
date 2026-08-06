#!/usr/bin/env python3
"""Build the technical closeout report for the EDR chat module."""

from __future__ import annotations

import argparse
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


NAVY = RGBColor(32, 55, 72)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(92, 105, 117)
GOLD = RGBColor(174, 121, 22)
WHITE = RGBColor(255, 255, 255)
INK = RGBColor(25, 36, 45)
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"
GREEN_FILL = "E7F4EC"
GOLD_FILL = "FFF4D6"
RED_FILL = "FCE8E6"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def set_run_font(run, size=11, color=INK, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, *, bold=False, color=INK, size=9.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(str(text))
    set_run_font(r, size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    set_run_font(run, size=9, color=MUTED)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def add_bottom_border(paragraph, color="D7DBE2", size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def make_numbering(doc, marker, num_format, left=540, hanging=271):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), num_format)
    lvl.append(num_fmt)
    r_pr = None
    if num_format == "bullet":
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), "Symbol")
        r_fonts.set(qn("w:hAnsi"), "Symbol")
        r_pr.append(r_fonts)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), marker)
    lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(left))
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(left))
    ind.set(qn("w:hanging"), str(hanging))
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    if r_pr is not None:
        lvl.append(r_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)
    p_pr.append(num_pr)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in specs.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def set_running_furniture(section):
    header = section.header
    p = header.paragraphs[0]
    p.text = ""
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("EDR DASHBOARD 1.1  |  INFORME TÉCNICO DE CHAT")
    set_run_font(r, size=8.5, color=MUTED, bold=True)
    add_bottom_border(p)

    footer = section.footer
    p = footer.paragraphs[0]
    p.text = ""
    p.paragraph_format.space_before = Pt(3)
    add_page_field(p)


def add_body(doc, text, *, bold_prefix=None, italic=False):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        a = p.add_run(bold_prefix)
        set_run_font(a, bold=True)
        b = p.add_run(text[len(bold_prefix):])
        set_run_font(b, italic=italic)
    else:
        r = p.add_run(text)
        set_run_font(r, italic=italic)
    return p


def add_bullet(doc, text, bullet_style):
    p = doc.add_paragraph(style=bullet_style)
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_number(doc, text, decimal_style):
    p = doc.add_paragraph(style=decimal_style)
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_callout(doc, title, body, apply_table_geometry, fill=LIGHT_FILL):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=10.5, color=NAVY, bold=True)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(body)
    set_run_font(r, size=10)
    apply_table_geometry(
        table,
        [TABLE_WIDTH_DXA],
        table_width_dxa=TABLE_WIDTH_DXA,
        indent_dxa=TABLE_INDENT_DXA,
        cell_margins_dxa=CELL_MARGINS,
    )
    prevent_row_split(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths, apply_table_geometry, *, header_fill=HEADER_FILL):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_table_header(header)
    prevent_row_split(header)
    for index, text in enumerate(headers):
        shade_cell(header.cells[index], header_fill)
        set_cell_text(header.cells[index], text, bold=True, color=NAVY, size=9)
    for row_values in rows:
        row = table.add_row()
        prevent_row_split(row)
        for index, value in enumerate(row_values):
            set_cell_text(row.cells[index], value, size=9)
    apply_table_geometry(
        table,
        widths,
        table_width_dxa=TABLE_WIDTH_DXA,
        indent_dxa=TABLE_INDENT_DXA,
        cell_margins_dxa=CELL_MARGINS,
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def build_report(output_path: Path, skill_root: Path):
    sys.path.insert(0, str(skill_root / "scripts"))
    from table_geometry import apply_table_geometry

    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])
    set_running_furniture(doc.sections[0])
    bullet_num_id = "List Bullet"
    decimal_num_id = "List Number"

    # Editorial cover — named override: 30 pt centered report title.
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(76)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    set_run_font(kicker.add_run("CIERRE TÉCNICO · MÓDULO /CHAT"), size=10, color=GOLD, bold=True)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    set_run_font(title.add_run("Mensajería profesional"), size=30, color=NAVY, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(5)
    set_run_font(
        subtitle.add_run("Rendimiento, presencia, experiencia tipo Teams y observabilidad"),
        size=15,
        color=DARK_BLUE,
    )
    descriptor = doc.add_paragraph()
    descriptor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    descriptor.paragraph_format.space_after = Pt(64)
    set_run_font(
        descriptor.add_run("Informe de implementación, métricas y evidencia de QA"),
        size=10.5,
        color=GOLD,
        bold=True,
    )
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(4)
    set_run_font(meta.add_run("Exploradores del Rey · EDR Dashboard 1.1"), size=12, color=NAVY, bold=True)
    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(meta2.add_run("Fecha de cierre: 6 de agosto de 2026"), size=9.5, color=MUTED, italic=True)

    doc.add_page_break()

    doc.add_heading("1. Resumen ejecutivo", level=1)
    add_body(
        doc,
        "El módulo /chat fue reforzado en tres frentes coordinados: eficiencia operativa, experiencia de mensajería y control de fallos. La lista de conversaciones ahora se pagina con cursor estable; la presencia se observa por lotes y por sesiones de dispositivo; y la interfaz incorpora borradores, búsqueda, menciones, navegación de no leídos y atajos sin alterar el diseño visual existente."
    )
    add_body(
        doc,
        "La comunicación en tiempo real también quedó endurecida: mensajes, reacciones, eliminaciones, recibos, escritura y resumen lateral comparten una ruta de sincronización inmediata. Los mensajes eliminados no conservan reacciones, y las reacciones iguales se agrupan por emoji con contador y detalle de participantes."
    )
    add_callout(
        doc,
        "Resultado de validación",
        "114 pruebas automatizadas aprobadas, compilación de producción completada, lint del alcance sin errores y cuatro escenarios de rendimiento medidos. La revisión del navegador confirmó ausencia de desbordamiento horizontal en 375, 768 y 1,440 píxeles; la sesión disponible no estaba autenticada, por lo que la validación visual interna del chat queda como prueba manual de liberación.",
        apply_table_geometry,
        fill=GREEN_FILL,
    )

    doc.add_heading("2. Alcance y arquitectura resultante", level=1)
    add_table(
        doc,
        ["Capa", "Cambio principal", "Efecto"],
        [
            ("API de chat", "Paginación por cursor, respuestas uniformes y trazas por solicitud.", "Menos lecturas, errores correlacionables y contrato predecible."),
            ("Firestore", "Índice compuesto por participante, eliminado, fecha e identificador.", "Orden estable y continuación sin saltos ni duplicados."),
            ("Tiempo real", "Listeners acotados y mezcla incremental de mensajes.", "Sincronización inmediata con memoria y costo controlados."),
            ("Presencia", "Sesiones por dispositivo, consultas en lotes de hasta 30 y latido reducido.", "Estado consistente entre pestañas y dispositivos."),
            ("Cliente", "SWR infinito, borradores locales, búsqueda combinada y accesibilidad ARIA.", "Navegación continua y recuperación de trabajo en progreso."),
            ("Seguridad", "Identidad autenticada, proyección pública, validación de adjuntos y logs saneados.", "Menor superficie de suplantación y fuga de datos."),
        ],
        [1800, 4050, 3510],
        apply_table_geometry,
    )

    doc.add_heading("3. Rendimiento y presencia", level=1)
    doc.add_heading("3.1 Paginación e índices definitivos", level=2)
    for item in [
        "La lista solicita 30 conversaciones y utiliza un registro centinela para determinar si existe una página siguiente.",
        "El cursor codifica fecha de actualización e identificador; ambos campos forman el orden definitivo y evitan ambigüedad cuando varias conversaciones comparten fecha.",
        "El resumen lateral utiliza últimoMensaje y omite la descarga de subcolecciones de mensajes y el enriquecimiento fotográfico por fila.",
        "La interfaz deduplica páginas por identificador y ofrece carga progresiva de conversaciones anteriores.",
    ]:
        add_bullet(doc, item, bullet_num_id)

    doc.add_heading("3.2 Reducción de lecturas y listeners", level=2)
    add_table(
        doc,
        ["Indicador", "Antes", "Después", "Mejora"],
        [
            ("Lecturas modeladas del listado (30 × 30 mensajes)", "901", "31", "−96.56 %"),
            ("Listeners de presencia para 65 contactos", "65", "3", "−95.38 %"),
            ("Escrituras de latido por minuto y sesión", "3", "1", "−66.67 %"),
            ("Autenticaciones subyacentes en 1,000 solicitudes concurrentes", "1,000 sin caché", "1", "Promesa en vuelo compartida"),
        ],
        [3660, 1600, 1600, 2500],
        apply_table_geometry,
    )
    citation = doc.add_paragraph()
    citation.paragraph_format.space_before = Pt(4)
    citation.paragraph_format.space_after = Pt(4)
    set_run_font(
        citation.add_run(
            "Base: benchmark local de Node v26.5.1 y modelo explícito de lecturas. Los porcentajes de lecturas y listeners son reducciones estructurales, no facturación observada de Firebase."
        ),
        size=8.5,
        color=MUTED,
        italic=True,
    )

    doc.add_heading("3.3 Presencia multidispositivo", level=2)
    add_body(
        doc,
        "Cada dispositivo mantiene su propia sesión de presencia. Una sesión visible vigente conserva el estado en línea aunque otra pestaña esté oculta; ocupado prevalece como estado manual; y la ausencia solo aparece cuando todas las sesiones activas dejan de estar visibles. Las sesiones vencidas no pueden sobrescribir una sesión reciente."
    )
    add_body(
        doc,
        "El latido se ejecuta cada 60 segundos y una sesión se considera vencida a los 150 segundos. Los listeners se liberan al desmontar el componente o cambiar el conjunto observado; el sidebar eleva una única suscripción por lote en lugar de crear una por contacto repetido."
    )

    doc.add_heading("4. Experiencia comparable con Teams", level=1)
    add_table(
        doc,
        ["Capacidad", "Implementación", "Validación"],
        [
            ("Borradores", "Persistencia local por miembro y conversación/destinatarios; limpieza tras envío exitoso.", "Prueba de aislamiento de claves."),
            ("Búsqueda", "Contactos, conversaciones, grupos y último mensaje dentro del directorio cargado.", "Coincidencia sin sensibilidad a mayúsculas ni acentos."),
            ("Atajos", "Ctrl/Cmd+K enfoca búsqueda; Alt+Shift+N recorre no leídos; Escape limpia búsqueda; Enter envía y Shift+Enter crea línea.", "Flujo de teclado y prueba de navegación circular."),
            ("Menciones", "Resolución de @nombre a identificadores únicos y sugerencias con roles listbox/option.", "Prueba con acentos y deduplicación."),
            ("Escritura", "Estado temporal con vencimiento local; las notificaciones obsoletas se retiran sin esperar otro snapshot.", "Prueba de expiración determinista."),
            ("Silencio", "Acción de conversación conservada y autorizada por participante.", "Contrato de acción en API."),
            ("No leídos", "Resumen global, actualización inmediata del tab Chats y salto circular al siguiente pendiente.", "Regresiones de resumen y entrega."),
            ("Responsive", "Navegación adaptable, controles compactos y entrada multilínea sin cambio de diseño.", "375/768/1,440 px sin scroll horizontal en shell disponible."),
            ("Accesibilidad", "Regiones log/status/search, aria-live, etiquetas de botones, lista de sugerencias y navegación por teclado.", "Lint del alcance y revisión DOM parcial."),
        ],
        [1550, 5210, 2600],
        apply_table_geometry,
    )
    add_callout(
        doc,
        "Alcance de búsqueda global",
        "La búsqueda combina todo el directorio de contactos con las páginas de conversaciones ya cargadas. Para búsqueda histórica de texto sobre todas las conversaciones será necesario un índice de búsqueda dedicado; no se amplió la lectura completa de Firestore para preservar las mejoras de rendimiento.",
        apply_table_geometry,
        fill=GOLD_FILL,
    )

    doc.add_heading("5. Mensajería, archivos y sincronización", level=1)
    for item in [
        "Los mensajes llegan por listener de cambios y se mezclan por identificador, conservando recibos locales y evitando duplicados.",
        "El indicador de escritura expira en el cliente y se limpia al enviar o abandonar, evitando el estado permanente «está escribiendo». ",
        "Reacciones, edición, eliminación y restauración actualizan el mensaje inmediatamente en ambos extremos.",
        "Un mensaje eliminado fuerza reacciones={} tanto al persistir como al recibir documentos históricos.",
        "Reacciones iguales aparecen una sola vez, muestran contador y exponen en una ventana los nombres de quienes reaccionaron.",
        "Los adjuntos admiten progreso, cancelación, reintento y limpieza recuperable; se validan ruta de conversación, MIME, tamaño y URL segura.",
        "Los recibos avanzan de enviado a entregado y visto de forma monótona entre pestañas y dispositivos.",
    ]:
        add_bullet(doc, item, bullet_num_id)

    doc.add_heading("6. Errores, observabilidad y protección de datos", level=1)
    add_table(
        doc,
        ["Control", "Comportamiento aplicado"],
        [
            ("Errores uniformes", "401 sesión, 403 autorización, errores de dominio y 500 genérico con código CHAT_INTERNAL_ERROR."),
            ("Cuota", "Mensaje único: «La fogata está encendiéndose. En unos minutos podrás continuar la conversación.» y Retry-After=60."),
            ("Correlación", "Cada solicitud registra requestId, método, endpoint, resultado y duración; el error devuelve X-Chat-Request-ID."),
            ("Logs saneados", "Solo código, estado, categoría, etapa y requestId. No se escriben cuerpos, correos, tokens ni mensajes internos."),
            ("Identidad", "El servidor deriva idMiembros del Bearer verificado y reemplaza cualquier remitente aportado por el cliente."),
            ("Datos públicos", "Contactos y participantes usan una proyección que excluye campos sensibles."),
            ("Adjuntos", "Storage limita operaciones a participantes y valida pertenencia de cada ruta a la conversación activa."),
        ],
        [2300, 7060],
        apply_table_geometry,
    )

    doc.add_heading("7. Métricas medidas", level=1)
    add_table(
        doc,
        ["Escenario", "Volumen", "Resultado", "Tiempo"],
        [
            ("Mezcla de cambios en tiempo real", "10,000 mensajes + 1,000 cambios", "10,000 mensajes, sin duplicados", "20.36 ms"),
            ("Construcción de página", "10,000 conversaciones", "30 elementos + cursor", "0.24 ms"),
            ("Presencia multidispositivo", "500 sesiones", "Estado online correcto", "0.94 ms"),
            ("Autenticación concurrente", "1,000 solicitudes", "1 lectura subyacente", "17.53 ms"),
        ],
        [3000, 2450, 2600, 1310],
        apply_table_geometry,
    )
    add_body(
        doc,
        "Estas mediciones son microbenchmarks ejecutados en el equipo de desarrollo el 6 de agosto de 2026. Sirven para comparar complejidad y detectar regresiones; no sustituyen una prueba end-to-end contra la latencia y cuota reales de Firebase en producción."
    )

    doc.add_heading("8. Matriz de QA final", level=1)
    add_table(
        doc,
        ["Área", "Evidencia", "Estado"],
        [
            ("Automatización", "114/114 pruebas Node aprobadas en 0.96 s.", "APROBADO"),
            ("Carga", "10k mensajes, 10k conversaciones, 500 sesiones y 1k solicitudes concurrentes.", "APROBADO"),
            ("Seguridad", "Bearer, identidad, permisos, proyección pública, adjuntos y redacción de logs.", "APROBADO"),
            ("Producción", "Next.js 16.0.10 compiló 258 páginas y rutas sin error.", "APROBADO"),
            ("Lint del alcance", "Archivos de chat, servidor y utilidades sin errores ni advertencias.", "APROBADO"),
            ("Responsive público", "375/768/1,440 px sin desbordamiento horizontal.", "APROBADO"),
            ("UI autenticada", "El navegador de QA fue redirigido al inicio de sesión; no se usaron credenciales.", "MANUAL REQUERIDO"),
            ("Navegadores/dispositivos reales", "Validar Chrome, Edge/Safari y Android/iOS con dos cuentas y dos dispositivos.", "MANUAL REQUERIDO"),
        ],
        [2200, 5310, 1850],
        apply_table_geometry,
    )

    doc.add_heading("9. Prueba manual rápida de liberación", level=1)
    steps = [
        "Abrir dos cuentas en navegadores o dispositivos distintos y confirmar que el estado en línea, ocupado y ausente se replica sin parpadeo.",
        "Enviar texto, mención, emoji y archivo; confirmar actualización inmediata del mensaje, tab Chats, contador de no leídos y recibos.",
        "Reaccionar con el mismo emoji desde dos usuarios; verificar un solo chip, contador 2 y detalle de participantes.",
        "Eliminar el mensaje reaccionado; verificar en ambos clientes que desaparecen cuerpo y reacciones en el mismo ciclo de listener.",
        "Escribir y detenerse sin enviar; confirmar que el indicador desaparece. Después redactar, recargar y verificar recuperación del borrador.",
        "Probar Ctrl/Cmd+K, Alt+Shift+N, Escape, Enter y Shift+Enter únicamente con teclado.",
        "Repetir en ancho móvil y escritorio, con zoom 200 %, lector de pantalla y foco visible.",
    ]
    for step in steps:
        add_number(doc, step, decimal_num_id)

    doc.add_heading("10. Despliegue y operación", level=1)
    add_callout(
        doc,
        "Acción obligatoria antes de producción",
        "Desplegar firestore.indexes.json en el proyecto correcto. La consulta paginada depende del índice compuesto; si el índice aún está construyéndose, Firestore devolverá un error controlado con requestId.",
        apply_table_geometry,
        fill=GOLD_FILL,
    )
    for item in [
        "Observar eventos chat_request por endpoint, resultado, durationMs, código y categoría; crear alertas de p95 y tasa de error.",
        "Supervisar RESOURCE_EXHAUSTED y la cuota de Firestore/Storage; el mensaje de fogata es recuperación de experiencia, no reemplaza ampliar o corregir capacidad.",
        "Ejecutar npm run test:chat-auth, lint acotado y npm run build en cada liberación.",
        "Repetir el benchmark después de cambios en mezcla de mensajes, paginación, autenticación o presencia.",
    ]:
        add_bullet(doc, item, bullet_num_id)

    doc.add_heading("11. Riesgos residuales y próximos pasos", level=1)
    add_table(
        doc,
        ["Riesgo / pendiente", "Tratamiento recomendado", "Prioridad"],
        [
            ("Búsqueda histórica total", "Incorporar servicio/índice dedicado si se requiere buscar el cuerpo de todos los mensajes sin cargar páginas.", "Media"),
            ("QA autenticado multinavegador", "Completar matriz manual con dos cuentas y dispositivos reales antes de liberar.", "Alta"),
            ("Cuota externa", "Medir lecturas/escrituras reales y ajustar plan, alertas o topología si reaparece RESOURCE_EXHAUSTED.", "Alta"),
            ("Accesibilidad integral", "Ejecutar auditoría con lector de pantalla, zoom 200 % y contraste en la sesión autenticada.", "Alta"),
            ("Telemetría agregada", "Conectar logs estructurados a panel y alertas; no añadir PII a etiquetas de métricas.", "Media"),
        ],
        [3000, 4860, 1500],
        apply_table_geometry,
    )

    doc.add_heading("12. Archivos técnicos principales", level=1)
    add_table(
        doc,
        ["Área", "Archivos"],
        [
            ("API y servidor", "src/app/api/chat/route.js; src/server/chat-pagination.mjs; chat-observability.mjs; chat-auth-core.mjs; chat-firestore-rest.mjs"),
            ("Cliente", "src/actions/chat.js; src/sections/chat/chat-nav.jsx; chat-message-input.jsx; chat-message-item.jsx; chat-message-list.jsx"),
            ("Tiempo real y presencia", "use-chat-realtime-sync.js; use-presence-status.js; realtime-sync.mjs; presence-state.mjs; src/lib/chat-presence.js"),
            ("Seguridad y archivos", "src/lib/axios.js; src/utils/firebase-file-storage.js; src/utils/chat-error.mjs; chat-reaction-core.mjs"),
            ("QA", "tests/chat/*.test.mjs; scripts/benchmark-chat.mjs; firestore.indexes.json"),
        ],
        [2200, 7160],
        apply_table_geometry,
    )

    add_body(
        doc,
        "Conclusión: el módulo queda con una base más cercana a una mensajería empresarial: sincronización inmediata, estados multidispositivo, costo acotado por página, controles de seguridad verificables y una ruta clara de operación. Los pendientes documentados son validaciones y capacidades de infraestructura adicionales, no fallos ocultados en la evidencia presentada.",
        italic=True,
    )

    # Ensure every section uses the same exact geometry and furniture.
    for section in doc.sections:
        configure_section(section)
        set_running_furniture(section)

    doc.core_properties.title = "Informe técnico de cierre del módulo Chat"
    doc.core_properties.subject = "Rendimiento, presencia, experiencia y QA"
    doc.core_properties.keywords = "chat, mensajería, Firestore, rendimiento, presencia, QA"
    doc.core_properties.comments = "Generado a partir de evidencia de implementación y pruebas."
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("output", type=Path)
    parser.add_argument("--skill-root", type=Path, required=True)
    args = parser.parse_args()
    build_report(args.output, args.skill_root)


if __name__ == "__main__":
    main()
